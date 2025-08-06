from pathlib import Path
import docx
from django.core.management.base import BaseCommand
from apps.workouts.models import WeeklyLesson, FinalVideo

DOCS = {
    "weekly": "Сценарии еженедельных  Видео.docx",
    "final": "финальные развернутые видео.docx",
}
ARCH_MAP = {"Наставник": "111", "Профессионал": "222", "Ровесник": "333"}


class Command(BaseCommand):
    help = "Импортирует weekly-lessons и финальные видео из Word"

    def handle(self, *a, **kw):
        root = Path("data/raw")
        created_weekly = 0
        updated_weekly = 0
        created_final = 0
        updated_final = 0
        
        # --- weekly ---
        weekly_path = root / DOCS["weekly"]
        if weekly_path.exists():
            doc = docx.Document(weekly_path)
            cur_week = None
            cur_arch = None
            script_buffer = []
            
            for p in doc.paragraphs:
                text = p.text.strip()
                
                # Detect week number
                if text.lower().startswith("неделя"):
                    # Save previous buffer if exists
                    if cur_week and cur_arch and script_buffer:
                        script = "\n".join(script_buffer)
                        obj, created = WeeklyLesson.objects.update_or_create(
                            week=cur_week, archetype=cur_arch, locale="ru",
                            defaults={"title": f"Неделя {cur_week} - {cur_arch}", "script": script},
                        )
                        if created:
                            created_weekly += 1
                        else:
                            updated_weekly += 1
                        script_buffer = []
                    
                    # Extract week number
                    try:
                        cur_week = int(text.split()[1])
                    except:
                        cur_week = None
                        
                # Detect archetype
                elif any(key in text for key in ARCH_MAP):
                    # Save previous buffer if exists
                    if cur_week and cur_arch and script_buffer:
                        script = "\n".join(script_buffer)
                        obj, created = WeeklyLesson.objects.update_or_create(
                            week=cur_week, archetype=cur_arch, locale="ru",
                            defaults={"title": f"Неделя {cur_week} - {cur_arch}", "script": script},
                        )
                        if created:
                            created_weekly += 1
                        else:
                            updated_weekly += 1
                        script_buffer = []
                    
                    cur_arch = ARCH_MAP[next(k for k in ARCH_MAP if k in text)]
                    
                # Collect script text
                elif text and cur_week and cur_arch:
                    script_buffer.append(text)
            
            # Save last buffer
            if cur_week and cur_arch and script_buffer:
                script = "\n".join(script_buffer)
                obj, created = WeeklyLesson.objects.update_or_create(
                    week=cur_week, archetype=cur_arch, locale="ru",
                    defaults={"title": f"Неделя {cur_week} - {cur_arch}", "script": script},
                )
                if created:
                    created_weekly += 1
                else:
                    updated_weekly += 1
        
        # --- final ---
        final_path = root / DOCS["final"]
        if final_path.exists():
            doc = docx.Document(final_path)
            cur_arch = None
            script_buffer = []
            
            for p in doc.paragraphs:
                text = p.text.strip()
                
                # Detect archetype
                if any(k in text for k in ARCH_MAP):
                    # Save previous buffer
                    if cur_arch and script_buffer:
                        script = "\n".join(script_buffer)
                        obj, created = FinalVideo.objects.update_or_create(
                            arch=cur_arch, locale="ru", 
                            defaults={"script": script}
                        )
                        if created:
                            created_final += 1
                        else:
                            updated_final += 1
                        script_buffer = []
                    
                    cur_arch = ARCH_MAP[next(k for k in ARCH_MAP if k in text)]
                    
                # Collect script text
                elif text and cur_arch:
                    script_buffer.append(text)
            
            # Save last buffer
            if cur_arch and script_buffer:
                script = "\n".join(script_buffer)
                obj, created = FinalVideo.objects.update_or_create(
                    arch=cur_arch, locale="ru", 
                    defaults={"script": script}
                )
                if created:
                    created_final += 1
                else:
                    updated_final += 1
        
        self.stdout.write(
            self.style.SUCCESS(
                f"Import done!\n"
                f"Weekly: Created {created_weekly}, Updated {updated_weekly}\n"
                f"Final: Created {created_final}, Updated {updated_final}"
            )
        )